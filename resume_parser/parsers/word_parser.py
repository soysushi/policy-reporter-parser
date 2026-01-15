"""
Word document parser implementation.

Design Decision: Using python-docx for .docx files
- Reliable extraction from modern Word documents
- Preserves paragraph structure
- Handles tables (common in resumes)
"""

import logging
from pathlib import Path

from docx import Document

from resume_parser.domain.exceptions import FileParsingError
from resume_parser.parsers.base import FileParser

logger = logging.getLogger(__name__)


class WordParser(FileParser):
    """
    Parser for Word (.docx) files.
    
    Design Decisions:
    - python-docx: Industry standard for .docx parsing
    - Extract paragraphs: Preserves document structure
    - Table support: Many resumes use tables for layout
    
    Scalability:
    - Stateless: Thread-safe, no shared state
    - Memory efficient: Lazy loading of document structure
    - Future: Support .doc (old format) via LibreOffice converter
    
    Limitations:
    - Only supports .docx (Office 2007+), not .doc (legacy)
    - Embedded images ignored (resumes rarely need OCR)
    """
    
    def parse(self, file_path: str) -> str:
        """
        Extract text from Word document.
        
        Strategy:
        1. Extract all paragraphs (main content)
        2. Extract table content (common in resumes)
        3. Merge with proper spacing
        
        N+1 Prevention: Single pass through document, build text once
        """
        path = Path(file_path)
        
        # Validate file exists and is readable
        if not path.exists():
            raise FileNotFoundError(f"Word file not found: {file_path}")
        
        if not path.is_file():
            raise FileParsingError(f"Path is not a file: {file_path}", file_path)
        
        try:
            return self._extract_text(file_path)
        except Exception as e:
            logger.error(f"Failed to parse Word document {file_path}: {e}")
            raise FileParsingError(
                f"Failed to parse Word document: {str(e)}",
                file_path
            ) from e
    
    def _extract_text(self, file_path: str) -> str:
        """
        Extract text from Word document including tables.
        
        Design Decision: Include tables for:
        - Many resumes use tables for contact info
        - Skills sections often in table format
        - Better coverage than paragraph-only extraction
        """
        doc = Document(file_path)
        text_parts = []
        
        # Extract paragraph text
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:  # Skip empty paragraphs
                text_parts.append(text)
        
        # Extract table text (common in resumes)
        for table in doc.tables:
            table_text = self._extract_table_text(table)
            if table_text:
                text_parts.append(table_text)
        
        if not text_parts:
            raise FileParsingError(
                "No text content found in Word document",
                file_path
            )
        
        # Join with newlines to preserve document structure
        return "\n".join(text_parts)
    
    def _extract_table_text(self, table) -> str:
        """
        Extract text from a table, preserving cell structure.
        
        Design Decision: Flatten table into text
        - Concatenate cells with spaces
        - Separate rows with newlines
        - Simple but effective for most resume tables
        """
        rows = []
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cells.append(cell_text)
            if cells:
                # Join cells with spaces to preserve info
                rows.append(" ".join(cells))
        
        return "\n".join(rows) if rows else ""
    
    def get_supported_extensions(self) -> list[str]:
        """
        Word parser supports .docx files.
        
        Design Decision: Only .docx (modern format)
        - .doc (legacy) requires different library (comtypes/win32com)
        - .docx covers 95%+ of modern resumes
        - Can add .doc support later if needed via conversion service
        """
        return ['.docx', '.DOCX']
